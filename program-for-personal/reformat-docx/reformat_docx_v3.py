import os
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from docx.oxml import OxmlElement

# 设置字体优先级列表
font_preferences = ["Segoe Print", "Ink Free", "微软雅黑"]

# 检查字体可用性（简单模拟）
def check_font_availability(font_name):
    return True  # 假设字体可用，实际可用系统特定方法检查

# 获取优先字体
def get_preferred_font():
    for font_name in font_preferences:
        if check_font_availability(font_name):
            return font_name
    return "Times New Roman"

# 调整文档边距和列
def adjust_margins_and_columns(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)

        cols = section._sectPr.xpath("./w:cols")
        if not cols:
            cols_element = OxmlElement('w:cols')
            cols_element.set(qn('w:num'), '2')
            cols_element.set(qn('w:equalWidth'), '1')
            section._sectPr.append(cols_element)
        else:
            cols[0].set(qn('w:num'), '2')
            cols[0].set(qn('w:equalWidth'), '1')

# 重新格式化文档
def reformat_docx(file_path, output_path):
    doc = Document(file_path)
    adjust_margins_and_columns(doc)
    
    total_length = sum(len(p.text) for p in doc.paragraphs)
    font_size = Pt(9.0) if total_length < 500 else Pt(9)
    preferred_font = get_preferred_font()

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = font_size
            run.font.name = preferred_font
            r = run._element
            rPr = r.rPr
            rFonts = rPr.rFonts if rPr.rFonts is not None else rPr.add_rFonts()
            rFonts.set(qn('w:eastAsia'), preferred_font)

    if not os.path.exists(output_path):
        os.makedirs(output_path)
    
    new_file_name = os.path.basename(file_path).replace(".docx", "_reformat.docx")
    new_file_path = os.path.join(output_path, new_file_name)
    doc.save(new_file_path)
    print(f"File saved: {new_file_path}")

# 扫描并重新格式化目录下的文件
def scan_and_reformat(source_path, output_path):
    for file_name in os.listdir(source_path):
        if file_name.endswith(".docx") and not file_name.endswith("_reformat.docx"):
            file_path = os.path.join(source_path, file_name)
            new_file_name = file_name.replace(".docx", "_reformat.docx")
            new_file_path = os.path.join(output_path, new_file_name)

            # 如果输出目录已经存在文件，则跳过
            if os.path.exists(new_file_path):
                print(f"Skipping {file_path}: Already reformatted.")
                continue

            print(f"Processing file: {file_path}")
            reformat_docx(file_path, output_path)

# 文件监控处理程序
class FileHandler(FileSystemEventHandler):
    def on_created(self, event):
        if event.is_directory:
            return
        if event.src_path.endswith(".docx") and not event.src_path.endswith("_reformat.docx"):
            file_name = os.path.basename(event.src_path)
            new_file_name = file_name.replace(".docx", "_reformat.docx")
            new_file_path = os.path.join(output_directory, new_file_name)

            # 如果输出目录已经存在文件，则跳过
            if os.path.exists(new_file_path):
                print(f"Skipping {event.src_path}: Already reformatted.")
                return

            print(f"New file detected: {event.src_path}")
            reformat_docx(event.src_path, output_directory)

def main():
    global output_directory
    source_directory = "/data/data/com.termux/files/home/EnjoyLibrary/1000h-english-learning-record/xiaolai_2024xxxx/docx-and-epub"
    output_directory = "/data/data/com.termux/files/home/EnjoyLibrary/1000h-english-learning-record/xiaolai_2024xxxx/reformat-docx"

    # 扫描现有的 .docx 文件并重新格式化
    scan_and_reformat(source_directory, output_directory)

    # 开始监视源文件目录和输出文件目录
    observer = Observer()
    event_handler = FileHandler()
    observer.schedule(event_handler, path=source_directory, recursive=True)
    observer.schedule(event_handler, path=output_directory, recursive=True)
    observer.start()

    try:
        print(f"Monitoring directory: {source_directory}")
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()

    observer.join()

if __name__ == "__main__":
    main()