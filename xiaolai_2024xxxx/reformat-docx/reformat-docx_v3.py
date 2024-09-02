from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import os
from docx.oxml import OxmlElement

# 设置字体优先级列表
font_preferences = ["Segoe Print", "Ink Free", "微软雅黑"]

def get_preferred_font():
    for font_name in font_preferences:
        if check_font_availability(font_name):
            return font_name
    return "Times New Roman"

def check_font_availability(font_name):
    # 实际实现字体检查功能
    return True

def adjust_margins_and_columns(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)

        # 确保存在 cols 元素，否则创建一个
        cols = section._sectPr.xpath("./w:cols")
        if not cols:
            cols_element = OxmlElement('w:cols')
            cols_element.set(qn('w:num'), '2')
            cols_element.set(qn('w:equalWidth'), '1')
            section._sectPr.append(cols_element)
        else:
            cols[0].set(qn('w:num'), '2')
            cols[0].set(qn('w:equalWidth'), '1')

def reformat_docx(file_path, output_path):
    doc = Document(file_path)
    adjust_margins_and_columns(doc)
    
    total_length = sum(len(p.text) for p in doc.paragraphs)
    
    if total_length < 500:
        font_size = Pt(9.0)
    elif total_length < 1000:
        font_size = Pt(9)
    else:
        font_size = Pt(9)

    preferred_font = get_preferred_font()

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = font_size
            run.font.name = preferred_font
            r = run._element
            rPr = r.rPr
            rFonts = rPr.rFonts if rPr.rFonts is not None else rPr.add_rFonts()
            rFonts.set(qn('w:eastAsia'), preferred_font)

    if not os.path.exists(output_path):
        os.makedirs(output_path)
    
    new_file_name = os.path.basename(file_path).replace(".docx", "_reformat.docx")
    new_file_path = os.path.join(output_path, new_file_name)
    doc.save(new_file_path)
    print(f"File saved: {new_file_path}")

def scan_and_reformat(source_path, output_path):
    for file_name in os.listdir(source_path):
        if file_name.endswith(".docx") and not file_name.endswith("_reformat.docx"):
            file_path = os.path.join(source_path, file_name)
            print(f"Processing file: {file_path}")
            reformat_docx(file_path, output_path)

if __name__ == "__main__":
    source_path = "/data/data/com.termux/files/home/EnjoyLibrary/1000h-english-learning-record/md-2-doc-epub/docx-and-epub"
    output_path = "/data/data/com.termux/files/home/EnjoyLibrary/1000h-english-learning-record/reformate-docx"
    scan_and_reformat(source_path, output_path)