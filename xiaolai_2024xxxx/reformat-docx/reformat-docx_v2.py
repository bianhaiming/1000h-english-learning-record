import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches

# 定义自定义路径
source_path = '/data/data/com.termux/files/home/EnjoyLibrary/1000h-english-learning-record/md-2-doc-epub/docx-and-epub'  # 源文件夹路径
output_path = '/data/data/com.termux/files/home/EnjoyLibrary/1000h-english-learning-record/md-2-doc-epub/docx-and-epub'  # 输出文件夹路径

def adjust_font_size(doc):
    max_pages = 1  # 限制为1页
    min_font_size = Pt(6)  # 最小字体大小
    default_font_size = Pt(8)  # 默认小六号字体
    tolerance = 0.1  # 容忍度（允许的微调幅度）

    while doc.sections[0].start_type != max_pages:
        # 设置段落字体大小为默认
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.size = default_font_size
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 保持左对齐
                run.font.name = 'SimSun'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        doc.save('temp.docx')
        temp_doc = Document('temp.docx')

        # 判断是否超过一页
        if len(temp_doc.element.xpath('//w:sectPr')) <= max_pages:
            break  # 字体大小已经合适
        else:
            # 将字体大小减少一点
            default_font_size -= Pt(tolerance)
            if default_font_size < min_font_size:
                break  # 已经达到最小字体，无法再调整

def reformat_docx(file_path, output_path):
    print(f"Processing file: {file_path}")  # 添加调试信息
    doc = Document(file_path)
    
    # 调整纸张大小为A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # 调整页边距
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # 添加文件名作为文档的第一段
    filename = os.path.basename(file_path).replace('.docx', '')
    first_para = doc.add_paragraph(filename)  # 正确插入段落
    if first_para:
        first_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
        first_para.runs[0].font.size = Pt(8)
        first_para.runs[0].font.name = 'SimSun'
        first_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        print(f"Added filename as the first paragraph: {filename}")  # 调试信息
    else:
        print(f"Failed to add filename as the first paragraph for file: {file_path}")

    # 添加空行
    doc.add_paragraph("")

    # 动态调整字体大小以适应一页
    adjust_font_size(doc)

    # 保存文件到输出路径，文件名加上_reformat后缀
    new_filename = filename + '_reformat.docx'
    output_file_path = os.path.join(output_path, new_filename)
    print(f"Saving file: {output_file_path}")  # 添加调试信息
    
    doc.save(output_file_path)
    print(f"File saved: {output_file_path}")  # 确认文件保存

def scan_and_reformat(source_path, output_path):
    # 获取输出文件夹中所有文件的文件名，便于去重
    output_files = [f.replace('_reformat.docx', '') for f in os.listdir(output_path) if f.endswith('_reformat.docx')]
    print(f"Output files: {output_files}")  # 添加调试信息
    
    # 扫描源文件夹，找到所有的.docx文件
    for root, dirs, files in os.walk(source_path):
        for file in files:
            if file.endswith('.docx'):
                file_name_without_ext = file.replace('.docx', '')
                # 检查是否已经处理过
                if file_name_without_ext not in output_files:
                    file_path = os.path.join(root, file)
                    reformat_docx(file_path, output_path)
                else:
                    print(f"File already processed: {file}")  # 已处理文件信息

# 执行文件扫描和重新排版
scan_and_reformat(source_path, output_path)
